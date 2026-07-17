#main.py
import os
from docx import Document
from src.redactor import PIIRedactor
from src.evaluator import evaluate_labeled_samples


def process_document(input_path: str, output_path: str, active_redactor: PIIRedactor):
    print(f"Loading document from {input_path}...")
    doc_obj = Document(input_path)

    print("Redacting paragraphs and tables...")
    for text_block in doc_obj.paragraphs:
        if text_block.text:
            text_block.text = active_redactor.redact_text(text_block.text)

    for doc_table in doc_obj.tables:
        for doc_row in doc_table.rows:
            for doc_cell in doc_row.cells:
                if doc_cell.text:
                    doc_cell.text = active_redactor.redact_text(doc_cell.text)

    doc_obj.save(output_path)
    print(f"Document processing complete! Saved to {output_path}")


def build_ground_truth_samples():
    """Real paragraphs/cells pulled from the prospectus, hand-labeled with
    (a) substrings that ARE PII and must be gone from the redacted output,
    and (b) substrings that are NOT PII and must survive untouched. Mixing
    both in every sample is what actually tests precision and recall
    together, instead of only recall (as the single synthetic sentence in
    the previous version did).
    """
    return [
        {
            "text": (
                "Contact Person: Sarthak Malvadkar, Company Secretary and Compliance "
                "Officer; Telephone: + 91 20 4505 3237; E-mail: cs.connect@kshinternational.com"
            ),
            "pii_terms": [
                "Sarthak Malvadkar", "+ 91 20 4505 3237", "cs.connect@kshinternational.com",
            ],
            "non_pii_terms": [
                "Company Secretary and Compliance Officer",
            ],
        },
        {
            "text": (
                "Kushal Subbayya Hegde, Pushpa Kushal Hegde, Rajesh Kushal Hegde, "
                "Rohit Kushal Hegde, Rakhi Girija Shetty ... are the Promoters of our Company."
            ),
            "pii_terms": [
                "Kushal Subbayya Hegde", "Pushpa Kushal Hegde", "Rajesh Kushal Hegde",
                "Rohit Kushal Hegde", "Rakhi Girija Shetty",
            ],
            "non_pii_terms": [
                "Promoters", "our Company",
            ],
        },
        {
            "text": (
                "Our Company was originally incorporated as \"Bhandary Metal Extrusion "
                "Private Limited\" under the provisions of the Companies Act, 1956, "
                "pursuant to a certificate of incorporation dated July 30, 1979."
            ),
            "pii_terms": [
                "Bhandary Metal Extrusion Private Limited",
            ],
            "non_pii_terms": [
                "Companies Act, 1956", "July 30, 1979", "our Company",
            ],
        },
        {
            "text": (
                "Bidders may contact the Company Secretary and Compliance Officer, "
                "Book Running Lead Managers or Registrar to the Offer in case of any "
                "pre-Offer or post-Offer related queries."
            ),
            "pii_terms": [],
            "non_pii_terms": [
                "Company Secretary and Compliance Officer", "Book Running Lead Managers",
                "Registrar to the Offer", "Bidders",
            ],
        },
        {
            "text": (
                "ICICI Securities Limited\nICICI Venture House, Appasaheb Marathe Marg, "
                "Prabhadevi, Mumbai 400025, Maharashtra, India\nTel: +91 22 6807 7100\n"
                "E-mail: ksh@icicisecurities.com\nSEBI Registration Number: INM000011179"
            ),
            "pii_terms": [
                "+91 22 6807 7100", "ksh@icicisecurities.com",
            ],
            "non_pii_terms": [
                "Maharashtra", "India", "SEBI Registration Number", "INM000011179",
                "ICICI Securities Limited",
            ],
        },
        {
            "text": (
                "(i) certificate of enrolment issued by Maharashtra Sales Tax Department "
                "under the Maharashtra State Tax on Professions, Trades, Callings and "
                "Employment Act, 1975."
            ),
            "pii_terms": [],
            "non_pii_terms": [
                "Maharashtra Sales Tax Department",
                "Maharashtra State Tax on Professions, Trades, Callings and Employment Act, 1975",
            ],
        },
    ]


def run_evaluation_on_document(active_redactor: PIIRedactor):
    print("\nRunning Evaluation Against Real Document Samples...")
    samples = build_ground_truth_samples()
    metrics, per_sample = evaluate_labeled_samples(
        active_redactor, samples, label="Prospectus Ground-Truth Sample"
    )

    print("Per-sample detail (for the README / eval report):")
    for i, result in enumerate(per_sample, start=1):
        print(f"\n--- Sample {i} ---")
        print("Original :", result["text"][:120].replace("\n", " "))
        print("Redacted :", result["redacted"][:120].replace("\n", " "))
        print("y_true   :", result["y_true"])
        print("y_pred   :", result["y_pred"])

    return metrics


if __name__ == "__main__":
    src_file = "data/Red Herring Prospectus.docx"
    dest_file = "data/redacted_output.docx"

    # NOTE: this list is now ADDED to PIIRedactor's built-in default_allow_list,
    # not a replacement for it (see the fix in redactor.py's __init__). Terms
    # here are extras specific to this document beyond the class's general
    # defaults.
    prospectus_ignore_set = {
        "SCSB", "SCSBs", "QIB", "QIBs", "RIIs",
        "Companies Act", "Companies Act, 2013", "Companies Act, 1956",
        "SEBI ICDR Regulations", "Depositories Act",
    }

    redactor_tool = PIIRedactor(custom_deny_list=prospectus_ignore_set)

    if not os.path.exists(src_file):
        print(f"Error: Place your input document at {src_file}")
    else:
        process_document(src_file, dest_file, redactor_tool)
        run_evaluation_on_document(redactor_tool)